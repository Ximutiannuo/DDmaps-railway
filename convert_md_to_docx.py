#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
将Markdown格式的俄语使用说明书转换为Word文档
"""
import sys
import os
import re

def markdown_to_docx(md_file_path, docx_file_path):
    """
    将Markdown文件转换为Word文档
    需要安装: pip install python-docx
    """
    try:
        from docx import Document  # type: ignore
        from docx.shared import Pt, RGBColor, Inches  # type: ignore
        from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore
        from docx.oxml.ns import qn  # type: ignore
    except ImportError:
        print("错误: 未安装 python-docx 库")
        print("请运行以下命令安装: pip install python-docx")
        return False
    
    # 读取Markdown文件
    try:
        with open(md_file_path, 'r', encoding='utf-8') as f:
            content = f.read()
    except FileNotFoundError:
        print(f"错误: 找不到文件 {md_file_path}")
        return False
    
    # 创建Word文档
    doc = Document()
    
    # 设置文档默认字体（支持俄语）
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    font._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    
    lines = content.split('\n')
    i = 0
    
    def clean_markdown(text):
        """清理Markdown标记"""
        # 移除链接，保留文本
        text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)
        # 移除代码块标记
        text = re.sub(r'`([^`]+)`', r'\1', text)
        # 处理粗体（保留内容，稍后处理）
        return text
    
    def add_formatted_text(paragraph, text):
        """添加格式化的文本，支持粗体"""
        # 查找粗体标记 **text**
        parts = re.split(r'(\*\*[^\*]+\*\*)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                # 粗体文本
                bold_text = part[2:-2]
                run = paragraph.add_run(bold_text)
                run.bold = True
            elif part:
                # 普通文本
                paragraph.add_run(part)
    
    while i < len(lines):
        line = lines[i].strip()
        
        if not line:
            # 空行
            doc.add_paragraph()
            i += 1
            continue
        
        # 处理标题
        if line.startswith('# '):
            # 一级标题
            title = line[2:].strip()
            title = clean_markdown(title)
            p = doc.add_heading(title, level=1)
            i += 1
        elif line.startswith('## '):
            # 二级标题
            title = line[3:].strip()
            title = clean_markdown(title)
            p = doc.add_heading(title, level=2)
            i += 1
        elif line.startswith('### '):
            # 三级标题
            title = line[4:].strip()
            title = clean_markdown(title)
            p = doc.add_heading(title, level=3)
            i += 1
        elif line.startswith('#### '):
            # 四级标题
            title = line[5:].strip()
            title = clean_markdown(title)
            p = doc.add_heading(title, level=4)
            i += 1
        elif line.startswith('---'):
            # 分隔线 - 跳过
            i += 1
        elif line.startswith('- ') or line.startswith('* '):
            # 列表项
            item_text = line[2:].strip()
            item_text = clean_markdown(item_text)
            p = doc.add_paragraph(style='List Bullet')
            add_formatted_text(p, item_text)
            i += 1
        elif re.match(r'^\d+\.\s', line):
            # 有序列表
            item_text = re.sub(r'^\d+\.\s+', '', line)
            item_text = clean_markdown(item_text)
            p = doc.add_paragraph(style='List Number')
            add_formatted_text(p, item_text)
            i += 1
        else:
            # 普通段落
            para_text = line
            para_text = clean_markdown(para_text)
            
            if para_text.strip():
                # 检查是否是Q&A格式
                if para_text.startswith('### Q') or para_text.startswith('### A'):
                    # Q&A标题
                    p = doc.add_heading(para_text.replace('### ', ''), level=3)
                elif para_text.startswith('**Q') or para_text.startswith('**A'):
                    # Q&A问题或答案
                    p = doc.add_paragraph()
                    add_formatted_text(p, para_text)
                else:
                    # 普通段落
                    p = doc.add_paragraph()
                    add_formatted_text(p, para_text)
            
            i += 1
    
    # 保存Word文档
    try:
        doc.save(docx_file_path)
        print(f"成功! Word文档已保存到: {docx_file_path}")
        return True
    except Exception as e:
        print(f"错误: 保存Word文档失败: {e}")
        import traceback
        traceback.print_exc()
        return False

if __name__ == '__main__':
    md_file = '司机端使用说明书_俄语版.md'
    docx_file = '司机端使用说明书_俄语版.docx'
    
    if len(sys.argv) > 1:
        md_file = sys.argv[1]
    if len(sys.argv) > 2:
        docx_file = sys.argv[2]
    
    print(f"正在将 {md_file} 转换为 {docx_file}...")
    print("请稍候...")
    success = markdown_to_docx(md_file, docx_file)
    
    if success:
        print("\n转换完成!")
        print(f"Word文档位置: {os.path.abspath(docx_file)}")
    else:
        print("\n转换失败，请检查错误信息")
